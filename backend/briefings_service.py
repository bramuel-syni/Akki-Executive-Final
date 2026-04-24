"""Briefings service — M12.

A Briefing is a 1–2 page document a NED or Executive can take into a meeting.
It bundles:
  • An opening paragraph written in AKKI's advisor voice (what should be front-of-mind)
  • For each included signal: the headline, evidence paragraph, and a sharp
    question to ask the board or management
  • Source citations preserved inline and as a footer

Exports to both PDF (via reportlab) and DOCX (via python-docx), with plain-text
fallback available as a third option.
"""
from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
)

AKKI_NAVY = colors.HexColor("#0A1F44")
AKKI_GOLD = colors.HexColor("#C9A961")
AKKI_INK = colors.HexColor("#1F2937")
AKKI_MUTED = colors.HexColor("#64748B")
AKKI_RULE = colors.HexColor("#E1E6ED")

TYPE_LABEL = {"risk": "Risk", "opportunity": "Opportunity", "gap": "Gap"}

# ---------------------------------------------------------------------------
# LLM prompt builder — takes signals + context object and returns a briefing
# ---------------------------------------------------------------------------
def build_briefing_prompt(
    *,
    context_name: str,
    role: str,
    context_object: Optional[Dict[str, Any]],
    signals: List[Dict[str, Any]],
    doc_ids_in_scope: List[str],
) -> str:
    co_answers = (context_object or {}).get("answers") or {}
    persona_bits = []
    for k in ("q1_role", "q3_focus_areas", "q5_prior_concerns", "q6_lens_preference", "q7_analytical_style"):
        v = co_answers.get(k)
        if v:
            persona_bits.append(f"  · {v}")
    persona_block = "\n".join(persona_bits) if persona_bits else "  · (generic board lens)"

    signals_block = []
    for i, s in enumerate(signals, 1):
        src_ids = [src.get("doc_id") for src in (s.get("sources") or [])]
        signals_block.append(
            f"Signal {i}\n"
            f"  id: {s.get('id')}\n"
            f"  type: {s.get('type', 'risk')} · confidence: {s.get('confidence', 'medium')}\n"
            f"  headline: {s.get('headline')}\n"
            f"  summary: {s.get('summary')}\n"
            f"  cited_doc_ids: {src_ids}\n"
        )
    signals_text = "\n".join(signals_block)

    return (
        f"You are preparing a private briefing for a {role} ahead of their next meeting "
        f"at «{context_name}». This briefing will be printed and taken into the room.\n\n"
        f"[WHO YOU'RE WRITING FOR]\n{persona_block}\n\n"
        f"[SIGNALS TO INCLUDE — these have already been generated and verified]\n{signals_text}\n\n"
        f"[VALID DOC IDS FOR CITATIONS]\n{doc_ids_in_scope}\n\n"
        f"Produce JSON with this shape:\n"
        f"{{\n"
        f'  "title": "short title, ≤ 70 chars, no quotes",\n'
        f'  "opening_paragraph": "3–5 sentences. Open with the single most important '
        f'thing they must not miss. Then frame the other items. Write in the advisor '
        f'voice — colleague not tool. No preamble, no filler.",\n'
        f'  "items": [\n'
        f'    {{\n'
        f'      "signal_id": "exact id string from the SIGNALS list above",\n'
        f'      "evidence": "2–3 sentence evidence paragraph — specific numbers from '
        f'the signal summary, rewritten for the briefing reader. May reference doc_ids '
        f'inline as [doc:xxx] using only doc_ids from the VALID list.",\n'
        f'      "question": "ONE sharp question the caller should ask in the meeting. '
        f'Declarative, not rhetorical. Uses specific numbers where possible. If the '
        f'caller is a CFO/Executive, phrase it as what they will say; if a NED/Chair, '
        f'phrase it as what they will ask management. ≤ 45 words."\n'
        f'    }}\n'
        f'  ],\n'
        f'  "closing_note": "Optional 1–2 sentence close. The one thing they should '
        f'feel in their gut as they walk in. May be null."\n'
        f"}}\n\n"
        f"CRITICAL: signal_id must be the exact UUID from the 'id:' line of one of "
        f"the signals above — copy it character for character. The items array should "
        f"contain one entry per signal you want in the briefing; you may reorder them "
        f"for narrative flow but do not invent new ones. Include all signals unless "
        f"two are genuinely redundant."
    )


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------
def _strip_citations_for_print(text: str) -> Tuple[str, List[str]]:
    """Replace [doc:xxx] with numeric footnotes and return cleaned text + id list (preserving order)."""
    ids_in_order: List[str] = []
    def repl(m):
        doc_id = m.group(1)
        if doc_id not in ids_in_order:
            ids_in_order.append(doc_id)
        idx = ids_in_order.index(doc_id) + 1
        return f"[{idx}]"
    cleaned = re.sub(r"\[doc:([a-f0-9-]+)\]", repl, text or "")
    return cleaned, ids_in_order


def render_pdf(briefing: Dict[str, Any], docs_by_id: Dict[str, Dict[str, Any]]) -> bytes:
    """Render a briefing to PDF bytes. `briefing` is the persisted briefing doc;
    `docs_by_id` maps doc_id → document (for the citations footer)."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=22 * mm, rightMargin=22 * mm,
        topMargin=22 * mm, bottomMargin=22 * mm,
        title=briefing.get("title", "AKKI Briefing"),
        author="AKKI",
    )

    styles = getSampleStyleSheet()
    style_overline = ParagraphStyle(
        "akki-overline", parent=styles["Normal"],
        textColor=AKKI_GOLD, fontName="Helvetica-Bold",
        fontSize=8, leading=10, spaceAfter=2,
        # fake letter-spacing via caps
    )
    style_title = ParagraphStyle(
        "akki-title", parent=styles["Normal"],
        textColor=AKKI_NAVY, fontName="Helvetica",
        fontSize=22, leading=26, spaceAfter=10,
    )
    style_meta = ParagraphStyle(
        "akki-meta", parent=styles["Normal"],
        textColor=AKKI_MUTED, fontName="Helvetica",
        fontSize=8.5, leading=11, spaceAfter=14,
    )
    style_opening = ParagraphStyle(
        "akki-opening", parent=styles["Normal"],
        textColor=AKKI_INK, fontName="Helvetica",
        fontSize=11, leading=17, spaceAfter=16,
        firstLineIndent=0,
    )
    style_item_header = ParagraphStyle(
        "akki-item-h", parent=styles["Normal"],
        textColor=AKKI_NAVY, fontName="Helvetica-Bold",
        fontSize=12, leading=16, spaceBefore=12, spaceAfter=4,
    )
    style_item_type = ParagraphStyle(
        "akki-item-type", parent=styles["Normal"],
        textColor=AKKI_GOLD, fontName="Helvetica-Bold",
        fontSize=7.5, leading=9, spaceAfter=2,
    )
    style_evidence = ParagraphStyle(
        "akki-ev", parent=styles["Normal"],
        textColor=AKKI_INK, fontName="Helvetica",
        fontSize=10, leading=15, spaceAfter=6,
    )
    style_question = ParagraphStyle(
        "akki-q", parent=styles["Normal"],
        textColor=AKKI_NAVY, fontName="Helvetica-Oblique",
        fontSize=10.5, leading=15, spaceAfter=8,
        leftIndent=10, borderPadding=0,
    )
    style_source = ParagraphStyle(
        "akki-src", parent=styles["Normal"],
        textColor=AKKI_MUTED, fontName="Helvetica",
        fontSize=8, leading=11, spaceAfter=2,
    )
    style_close = ParagraphStyle(
        "akki-close", parent=styles["Normal"],
        textColor=AKKI_INK, fontName="Helvetica-Oblique",
        fontSize=10.5, leading=16, spaceBefore=12, spaceAfter=6,
    )

    flow = []
    # Header
    flow.append(Paragraph("PRIVATE · AKKI BRIEFING", style_overline))
    flow.append(Paragraph(_escape(briefing.get("title", "Briefing")), style_title))
    created_at = briefing.get("created_at", "")
    meta_bits = []
    if briefing.get("context_name"):
        meta_bits.append(briefing["context_name"])
    if created_at:
        try:
            dt = datetime.fromisoformat(created_at)
            meta_bits.append(dt.strftime("%-d %B %Y · %H:%M"))
        except Exception:
            meta_bits.append(created_at[:10])
    meta_bits.append(f"v{briefing.get('version', 1)}")
    meta_bits.append(f"{len(briefing.get('items', []))} items")
    flow.append(Paragraph(" · ".join(meta_bits), style_meta))

    # Rule
    flow.append(_hr())

    # Collect all citations across opening + items so footer is stable
    all_ids: List[str] = []
    def _collect(text: str):
        for m in re.finditer(r"\[doc:([a-f0-9-]+)\]", text or ""):
            d = m.group(1)
            if d not in all_ids and d in docs_by_id:
                all_ids.append(d)

    _collect(briefing.get("opening_paragraph", ""))
    for it in briefing.get("items", []):
        _collect(it.get("evidence", ""))
        _collect(it.get("question", ""))

    def _with_footnotes(text: str) -> str:
        def repl(m):
            d = m.group(1)
            if d not in all_ids:
                return ""
            idx = all_ids.index(d) + 1
            return f'<font color="#C9A961"><b>[{idx}]</b></font>'
        return re.sub(r"\[doc:([a-f0-9-]+)\]", repl, _escape(text))

    # Opening
    opening = briefing.get("opening_paragraph") or ""
    if opening:
        flow.append(Paragraph(_with_footnotes(opening), style_opening))

    # Items
    for i, it in enumerate(briefing.get("items", []), 1):
        sig_type = (it.get("signal_type") or "risk").upper()
        conf = (it.get("confidence") or "medium").upper()
        flow.append(Paragraph(f"{i:02d} · {sig_type} · {conf} CONFIDENCE", style_item_type))
        flow.append(Paragraph(_escape(it.get("signal_headline") or "Untitled"), style_item_header))
        ev = it.get("evidence") or ""
        if ev:
            flow.append(Paragraph(_with_footnotes(ev), style_evidence))
        q = it.get("question") or ""
        if q:
            flow.append(Paragraph(f'<font color="#C9A961">→</font> &nbsp;<i>{_with_footnotes(q)}</i>', style_question))

    # Closing
    closing = briefing.get("closing_note")
    if closing:
        flow.append(_hr())
        flow.append(Paragraph(_escape(closing), style_close))

    # Citations footer
    if all_ids:
        flow.append(Spacer(1, 14))
        flow.append(_hr())
        flow.append(Paragraph("SOURCES", style_overline))
        for idx, did in enumerate(all_ids, 1):
            d = docs_by_id.get(did, {})
            name = d.get("name") or did[:8]
            trust = d.get("data_trust") or "unrated"
            flow.append(Paragraph(f"<b>[{idx}]</b> {_escape(name)} · data trust: {_escape(trust)}", style_source))

    flow.append(Spacer(1, 8))
    flow.append(_hr())
    flow.append(Paragraph("AKKI · grounded briefings · private to the caller", style_source))

    doc.build(flow)
    return buf.getvalue()


def _escape(text: str) -> str:
    """reportlab uses XML-like parsing in Paragraph; escape the basics."""
    if not text:
        return ""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _hr():
    return Table([[""]], colWidths=[None], rowHeights=[0.4 * mm], style=TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), AKKI_RULE),
        ("LINEABOVE", (0, 0), (-1, -1), 0.4, AKKI_RULE),
    ]))


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------
def render_docx(briefing: Dict[str, Any], docs_by_id: Dict[str, Dict[str, Any]]) -> bytes:
    doc = Document()

    # Slim margins + baseline style
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # Header overline
    p = doc.add_paragraph()
    r = p.add_run("PRIVATE · AKKI BRIEFING")
    r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0xC9, 0xA9, 0x61)

    # Title
    p = doc.add_paragraph()
    r = p.add_run(briefing.get("title", "Briefing"))
    r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x0A, 0x1F, 0x44)

    # Meta
    created_at = briefing.get("created_at", "")
    meta_bits = []
    if briefing.get("context_name"): meta_bits.append(briefing["context_name"])
    if created_at:
        try:
            dt = datetime.fromisoformat(created_at)
            meta_bits.append(dt.strftime("%-d %B %Y · %H:%M"))
        except Exception:
            meta_bits.append(created_at[:10])
    meta_bits.append(f"v{briefing.get('version', 1)}")
    meta_bits.append(f"{len(briefing.get('items', []))} items")
    p = doc.add_paragraph()
    r = p.add_run(" · ".join(meta_bits))
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Collect citations for footnote numbering
    all_ids: List[str] = []
    def _collect(text: str):
        for m in re.finditer(r"\[doc:([a-f0-9-]+)\]", text or ""):
            d = m.group(1)
            if d not in all_ids and d in docs_by_id:
                all_ids.append(d)
    _collect(briefing.get("opening_paragraph", ""))
    for it in briefing.get("items", []):
        _collect(it.get("evidence", ""))
        _collect(it.get("question", ""))

    def _render_runs(paragraph, text: str, *, italic: bool = False, base_color: RGBColor = None):
        # Splits on [doc:xxx] and renders footnote numbers in gold
        parts = re.split(r"(\[doc:[a-f0-9-]+\])", text or "")
        for part in parts:
            m = re.match(r"^\[doc:([a-f0-9-]+)\]$", part)
            if m:
                d = m.group(1)
                if d in all_ids:
                    run = paragraph.add_run(f"[{all_ids.index(d) + 1}]")
                    run.font.size = Pt(9); run.bold = True
                    run.font.color.rgb = RGBColor(0xC9, 0xA9, 0x61)
            else:
                if part:
                    run = paragraph.add_run(part)
                    if italic: run.italic = True
                    if base_color: run.font.color.rgb = base_color

    doc.add_paragraph().add_run("─" * 80).font.color.rgb = RGBColor(0xE1, 0xE6, 0xED)

    opening = briefing.get("opening_paragraph") or ""
    if opening:
        p = doc.add_paragraph()
        _render_runs(p, opening, base_color=RGBColor(0x1F, 0x29, 0x37))

    # Items
    for i, it in enumerate(briefing.get("items", []), 1):
        # Type + confidence overline
        p = doc.add_paragraph()
        r = p.add_run(f"{i:02d} · {(it.get('signal_type') or 'risk').upper()} · {(it.get('confidence') or 'medium').upper()} CONFIDENCE")
        r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0xC9, 0xA9, 0x61)
        # Headline
        p = doc.add_paragraph()
        r = p.add_run(it.get("signal_headline") or "Untitled")
        r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x0A, 0x1F, 0x44)
        # Evidence
        if it.get("evidence"):
            p = doc.add_paragraph()
            _render_runs(p, it["evidence"], base_color=RGBColor(0x1F, 0x29, 0x37))
        # Question
        if it.get("question"):
            p = doc.add_paragraph()
            arrow = p.add_run("→  ")
            arrow.bold = True; arrow.font.color.rgb = RGBColor(0xC9, 0xA9, 0x61)
            _render_runs(p, it["question"], italic=True, base_color=RGBColor(0x0A, 0x1F, 0x44))

    # Closing
    if briefing.get("closing_note"):
        doc.add_paragraph().add_run("─" * 80).font.color.rgb = RGBColor(0xE1, 0xE6, 0xED)
        p = doc.add_paragraph()
        r = p.add_run(briefing["closing_note"])
        r.italic = True; r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # Sources
    if all_ids:
        doc.add_paragraph().add_run("─" * 80).font.color.rgb = RGBColor(0xE1, 0xE6, 0xED)
        p = doc.add_paragraph()
        r = p.add_run("SOURCES")
        r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0xC9, 0xA9, 0x61)
        for idx, did in enumerate(all_ids, 1):
            d = docs_by_id.get(did, {})
            name = d.get("name") or did[:8]
            trust = d.get("data_trust") or "unrated"
            p = doc.add_paragraph()
            r = p.add_run(f"[{idx}] ")
            r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xC9, 0xA9, 0x61)
            r2 = p.add_run(f"{name} · data trust: {trust}")
            r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Board deck PDF — landscape A4 slide-per-signal, designed to be projected or
# printed and walked into a board room. The visual system stays aligned with
# the cream/oxblood editorial brief: deep navy headings, oxblood accents,
# muted slate body, generous margins. One logical slide per page.
# ---------------------------------------------------------------------------
AKKI_CREAM_BG = colors.HexColor("#F7F3EA")
AKKI_OXBLOOD = colors.HexColor("#8B2E2B")
AKKI_OXBLOOD_SOFT = colors.HexColor("#C58A88")


def _slide_header(footer_text: str, page_num: int, total: int):
    """Top-of-slide overline + a thin oxblood rule + page counter."""
    return [
        Paragraph(
            f'<font color="#8B2E2B">●</font>  <font color="#64748B">'
            f'{footer_text}  ·  slide {page_num} / {total}</font>',
            ParagraphStyle(
                "slide-overline", fontName="Helvetica-Bold",
                fontSize=7.5, leading=10, textColor=AKKI_MUTED,
                spaceAfter=4,
            ),
        ),
        Table([[""]], colWidths=[None], rowHeights=[0.6 * mm], style=TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), AKKI_OXBLOOD),
        ])),
        Spacer(1, 14),
    ]


def render_board_deck_pdf(briefing: Dict[str, Any], docs_by_id: Dict[str, Dict[str, Any]]) -> bytes:
    """Landscape slide-deck PDF — one signal per page, built to be presented.

    Slide sequence:
      1. Cover — title, context, date
      2. Executive summary — opening paragraph
      3..N. One slide per item (headline, evidence, sharpest question)
      N+1. Closing question / next steps (if set)
      Last. Sources list
    """
    buf = io.BytesIO()
    pagesize = landscape(A4)
    doc = SimpleDocTemplate(
        buf, pagesize=pagesize,
        leftMargin=28 * mm, rightMargin=28 * mm,
        topMargin=22 * mm, bottomMargin=18 * mm,
        title=briefing.get("title", "AKKI Board Deck"),
        author="AKKI",
    )

    styles = getSampleStyleSheet()
    style_overline = ParagraphStyle(
        "deck-overline", parent=styles["Normal"],
        textColor=AKKI_OXBLOOD, fontName="Helvetica-Bold",
        fontSize=9, leading=11, spaceAfter=4,
    )
    style_cover_title = ParagraphStyle(
        "deck-cover-title", parent=styles["Normal"],
        textColor=AKKI_INK, fontName="Times-Roman",
        fontSize=42, leading=48, spaceAfter=12,
    )
    style_cover_meta = ParagraphStyle(
        "deck-cover-meta", parent=styles["Normal"],
        textColor=AKKI_MUTED, fontName="Helvetica",
        fontSize=11, leading=15, spaceAfter=6,
    )
    style_slide_kicker = ParagraphStyle(
        "deck-kicker", parent=styles["Normal"],
        textColor=AKKI_OXBLOOD, fontName="Helvetica-Bold",
        fontSize=9, leading=11, spaceAfter=6,
    )
    style_slide_title = ParagraphStyle(
        "deck-slide-title", parent=styles["Normal"],
        textColor=AKKI_INK, fontName="Times-Roman",
        fontSize=26, leading=32, spaceAfter=12,
    )
    style_slide_body = ParagraphStyle(
        "deck-slide-body", parent=styles["Normal"],
        textColor=AKKI_INK, fontName="Helvetica",
        fontSize=13, leading=20, spaceAfter=10,
    )
    style_question = ParagraphStyle(
        "deck-question", parent=styles["Normal"],
        textColor=AKKI_OXBLOOD, fontName="Times-Italic",
        fontSize=14, leading=20, spaceBefore=10, spaceAfter=6,
        leftIndent=12,
    )
    style_source_line = ParagraphStyle(
        "deck-src", parent=styles["Normal"],
        textColor=AKKI_MUTED, fontName="Helvetica",
        fontSize=10, leading=14, spaceAfter=4,
    )

    # Collect citations up-front so each slide can footnote them
    all_ids: List[str] = []

    def _collect(text: str):
        for m in re.finditer(r"\[doc:([a-f0-9-]+)\]", text or ""):
            d = m.group(1)
            if d not in all_ids and d in docs_by_id:
                all_ids.append(d)

    _collect(briefing.get("opening_paragraph", ""))
    for it in briefing.get("items", []):
        _collect(it.get("evidence", ""))
        _collect(it.get("question", ""))

    def _with_footnotes(text: str) -> str:
        def repl(m):
            d = m.group(1)
            if d not in all_ids:
                return ""
            idx = all_ids.index(d) + 1
            return f'<font color="#8B2E2B"><b>[{idx}]</b></font>'
        return re.sub(r"\[doc:([a-f0-9-]+)\]", repl, _escape(text))

    # Compute total pages for slide counter
    item_count = len(briefing.get("items", []))
    total_slides = 2 + item_count + (1 if briefing.get("closing_note") else 0) + (1 if all_ids else 0)

    flow: List[Any] = []

    context_name = briefing.get("context_name") or "Board"

    # ── Slide 1 · Cover ────────────────────────────────────────────────────
    flow.extend(_slide_header(f"{context_name} · PRIVATE BOARD DECK", 1, total_slides))
    flow.append(Spacer(1, 40))
    flow.append(Paragraph("PRESENTED TO THE BOARD", style_overline))
    flow.append(Paragraph(_escape(briefing.get("title", "Board Briefing")), style_cover_title))
    created_at = briefing.get("created_at", "")
    meta_line = context_name
    if created_at:
        try:
            dt = datetime.fromisoformat(created_at)
            meta_line += f"  ·  {dt.strftime('%-d %B %Y')}"
        except Exception:
            pass
    meta_line += f"  ·  v{briefing.get('version', 1)}  ·  {item_count} item{'s' if item_count != 1 else ''}"
    flow.append(Paragraph(meta_line, style_cover_meta))
    flow.append(Spacer(1, 20))
    flow.append(Paragraph(
        '<font color="#8B2E2B">❝</font>  <i>Read sharper. Challenge better. '
        'Decide with receipts.</i>  <font color="#8B2E2B">❞</font>',
        ParagraphStyle("cover-tag", fontName="Times-Italic", fontSize=12,
                       textColor=AKKI_MUTED, leading=18),
    ))
    flow.append(PageBreak())

    # ── Slide 2 · Executive summary ────────────────────────────────────────
    flow.extend(_slide_header(f"{context_name} · PRIVATE BOARD DECK", 2, total_slides))
    flow.append(Paragraph("EXECUTIVE SUMMARY", style_slide_kicker))
    flow.append(Paragraph("What the board should front-of-mind", style_slide_title))
    opening = briefing.get("opening_paragraph") or ""
    if opening:
        flow.append(Paragraph(_with_footnotes(opening), style_slide_body))
    else:
        flow.append(Paragraph(
            f"The briefing contains {item_count} item{'s' if item_count != 1 else ''} "
            f"across the board's scope. Each item follows on the next slide — "
            f"headline, evidence, and the sharpest question to raise.",
            style_slide_body,
        ))
    flow.append(PageBreak())

    # ── Slides 3..N · One per signal ───────────────────────────────────────
    for idx, it in enumerate(briefing.get("items", []), 1):
        slide_num = 2 + idx
        flow.extend(_slide_header(f"{context_name} · PRIVATE BOARD DECK", slide_num, total_slides))
        sig_type = (it.get("signal_type") or "signal").upper()
        conf = (it.get("confidence") or "medium").upper()
        flow.append(Paragraph(
            f"ITEM {idx:02d}  ·  {sig_type}  ·  {conf} CONFIDENCE",
            style_slide_kicker,
        ))
        flow.append(Paragraph(_escape(it.get("signal_headline") or "Untitled"), style_slide_title))
        ev = it.get("evidence") or ""
        if ev:
            flow.append(Paragraph(_with_footnotes(ev), style_slide_body))
        q = it.get("question") or ""
        if q:
            flow.append(Paragraph(
                f'<font color="#8B2E2B">→</font> &nbsp; Question for the chair: <i>{_with_footnotes(q)}</i>',
                style_question,
            ))
        # Per-item source chips (small, muted)
        srcs = it.get("sources") or []
        if srcs:
            chip_bits = []
            for s in srcs[:4]:
                nm = (s.get("doc_name") or "source")[:44]
                trust = s.get("data_trust") or "mixed"
                chip_bits.append(f'<font color="#8B2E2B">●</font> {_escape(nm)} <font color="#94A3B8">· {trust}</font>')
            flow.append(Spacer(1, 8))
            flow.append(Paragraph(
                "  &nbsp;&nbsp;".join(chip_bits),
                ParagraphStyle("deck-chips", fontName="Helvetica", fontSize=9,
                               textColor=AKKI_MUTED, leading=13),
            ))
        flow.append(PageBreak())

    # ── Closing slide (if set) ─────────────────────────────────────────────
    if briefing.get("closing_note"):
        slide_num = 3 + item_count
        flow.extend(_slide_header(f"{context_name} · PRIVATE BOARD DECK", slide_num, total_slides))
        flow.append(Paragraph("IN CLOSING", style_slide_kicker))
        flow.append(Paragraph("One thing to take away", style_slide_title))
        flow.append(Paragraph(_escape(briefing["closing_note"]), style_slide_body))
        flow.append(PageBreak())

    # ── Sources slide ──────────────────────────────────────────────────────
    if all_ids:
        flow.extend(_slide_header(f"{context_name} · PRIVATE BOARD DECK", total_slides, total_slides))
        flow.append(Paragraph("RECEIPTS", style_slide_kicker))
        flow.append(Paragraph("Every claim, to the document it came from", style_slide_title))
        for idx, did in enumerate(all_ids, 1):
            d = docs_by_id.get(did, {})
            name = d.get("name") or did[:8]
            trust = d.get("data_trust") or "unrated"
            flow.append(Paragraph(
                f'<font color="#8B2E2B"><b>[{idx}]</b></font> &nbsp; '
                f'{_escape(name)} <font color="#94A3B8">· data trust: {_escape(trust)}</font>',
                style_source_line,
            ))

    doc.build(flow)
    return buf.getvalue()
