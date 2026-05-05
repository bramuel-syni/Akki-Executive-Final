"""Solva v3 — artefact PDF + DOCX export.

Phase I.4 of the Solva UX redesign.

Two public functions:
    * ``build_pdf(session)``  -> bytes  (WeasyPrint)
    * ``build_docx(session)`` -> bytes  (python-docx, BytesIO)

The shape of ``session`` is the dict stored in
``db.solva_v2_sessions`` (see ``backend/routers/solva_v2.py``). The
artefact composition is described in §5 of the Solva UX Redesign Brief.

Refusal sessions (status in {"refused", "blocked_hard", "blocked_soft"}
or ``synthesis`` missing) get the four-section refusal artefact instead
of the standard five-section one (brief §5.5).

Design notes:
  * Probability bars are rendered as CSS divs in HTML/PDF (no images,
    so the PDF stays small and crisp at any zoom). In DOCX they become
    a 1-row inline table, per the brief.
  * Same Georgia / Calibri palette as the web. Confidence intervals are
    *derived* from the band: the engine stores ``confidence_pct`` and a
    ``confidence_band``; the bar's CI extension uses a band-derived
    width because the engine does not (yet) emit per-claim low/high.
"""
from __future__ import annotations

import io
import re
import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from jinja2 import Environment, FileSystemLoader, select_autoescape

logger = logging.getLogger("akki.solva.export")

_TEMPLATE_DIR = Path(__file__).resolve().parent / "templates"
_env = Environment(
    loader=FileSystemLoader(str(_TEMPLATE_DIR)),
    autoescape=select_autoescape(["html"]),
)

SUBMODULE_LABELS = {
    "seek_clarity":        "Seek Clarity",
    "develop_strategy":    "Develop Strategy",
    "simulate_hypothesis": "Simulate Hypothesis",
    "get_perspective":     "See Different Perspectives",
}

# CI half-width derived from the band — the engine does not (yet) store
# per-claim low/high. The width is conservative: lower-conviction claims
# get wider intervals, higher-conviction ones tighter.
_BAND_HALF_WIDTH = {
    "Unlikely":         15,
    "Possible":         10,
    "Likely":            8,
    "High-conviction":   5,
}

_TIER_MARKER_RE = re.compile(r"\[T:[a-zA-Z_]+\]")
_REC_HEADING_RE = re.compile(r"^\s*Recommendation\s*(\d+)\s*:\s*", re.IGNORECASE)


# ---------------------------------------------------------------------------
# Shared shaping
# ---------------------------------------------------------------------------
def _strip_tier_markers(text: str) -> str:
    return _TIER_MARKER_RE.sub("", text or "").replace("  ", " ").strip()


def _split_paragraphs(body: str) -> List[str]:
    if not body:
        return []
    cleaned = _strip_tier_markers(body)
    paras = [p.strip() for p in re.split(r"\n\s*\n", cleaned) if p.strip()]
    return paras or [cleaned]


def _is_refusal(session: Dict[str, Any]) -> bool:
    """Return True when the session should render the refusal artefact.

    Refusal triggers:
      * session.status in the locked-out set
      * synthesis missing entirely (engine bailed before producing one)
      * synthesis.body empty and reasoning_audit_log carries a refusal
        decision (keeps the brief's promise that refusal is consequential).
    """
    status = (session.get("status") or "").lower()
    if status in {"refused", "blocked_hard", "blocked_soft"}:
        return True
    synth = session.get("synthesis") or {}
    if not synth.get("body"):
        for entry in (session.get("reasoning_audit_log") or []):
            if (entry.get("engine") or "").lower() == "refusal":
                out = entry.get("output") or {}
                if out.get("verdict") in {"refuse", "soft_block", "hard_block"}:
                    return True
    return False


def _scenarios_from_claims(claims: List[Dict[str, Any]], cap: int = 5) -> List[Dict[str, Any]]:
    """Project parsed claims into scenario rows for the bar visual.

    Each claim becomes a scenario row with:
        label: short heading (first 6-8 words of the sentence)
        desc:  the rest of the sentence (or empty)
        pct:   confidence_pct  (engine value, never null after weighting)
        low:   pct - half_width  (clamped to 0)
        high:  pct + half_width  (clamped to 100)
    """
    out: List[Dict[str, Any]] = []
    if not claims:
        return out
    for c in claims[:cap]:
        if not isinstance(c, dict):
            continue
        text = _strip_tier_markers(c.get("text") or "")
        if not text:
            continue
        words = text.split()
        head = " ".join(words[:8])
        if len(words) > 8:
            # Keep punctuation off the head if possible.
            head = head.rstrip(",.;:") + "…"
        rest = " ".join(words[8:]) if len(words) > 8 else ""
        try:
            pct = int(c.get("confidence_pct") or 50)
        except (TypeError, ValueError):
            pct = 50
        pct = max(0, min(100, pct))
        band = c.get("confidence_band") or "Possible"
        half = _BAND_HALF_WIDTH.get(band, 10)
        low = max(0, pct - half)
        high = min(100, pct + half)
        out.append({
            "label": head or "Scenario",
            "desc":  rest,
            "pct":   pct,
            "low":   low,
            "high":  high,
            "band":  band,
            "tier":  c.get("tier") or "",
        })
    # Sort by pct descending so the strongest scenario reads first.
    out.sort(key=lambda r: -r["pct"])
    return out


def _recommendations_for_artefact(session: Dict[str, Any]) -> List[Dict[str, str]]:
    synth = session.get("synthesis") or {}
    raw = synth.get("recommendations") or []
    out: List[Dict[str, str]] = []
    for r in raw:
        if isinstance(r, dict):
            heading = r.get("heading") or r.get("title") or ""
            body = r.get("body") or r.get("text") or ""
        else:
            text = str(r)
            m = _REC_HEADING_RE.match(text)
            if m:
                heading = f"Recommendation {m.group(1)}"
                body = text[m.end():].strip()
            else:
                heading = "Recommendation"
                body = text
        out.append({
            "heading": _strip_tier_markers(heading)[:80],
            "body":    _strip_tier_markers(body),
        })
    return out


def _sensitivity_items(session: Dict[str, Any], scenarios: List[Dict[str, Any]]) -> List[str]:
    """Sensitivity drivers — the 2-3 inputs that would most shift the read.

    Source preference order:
      1. ``session._sensitivity_drivers`` if the engine ever populates it
      2. ``synthesis.sensitivity[]`` if the orchestrator splits it out
      3. derived from the lowest-tier (most assumption-heavy) claims —
         these are the levers an executive should pressure-test
    """
    drivers = session.get("_sensitivity_drivers")
    if drivers:
        return [str(d) for d in drivers if d][:3]
    synth = session.get("synthesis") or {}
    if synth.get("sensitivity"):
        return [str(s) for s in synth["sensitivity"] if s][:3]

    # Derived: surface 1-2 scenarios whose tier is `domain_prior`,
    # `user_assertion`, or `speculation` — those are the assumption-heavy
    # claims that, if proved wrong, would shift the diagnosis most.
    weak_tiers = {"domain_prior", "user_assertion", "speculation"}
    candidates = [s for s in scenarios if s.get("tier") in weak_tiers]
    items: List[str] = []
    for s in candidates[:3]:
        line = (
            f"If the assumption that {s['label'].lower()} no longer holds, "
            f"the {s['pct']}% read would move materially."
        )
        items.append(line)
    return items


def _tension_items(session: Dict[str, Any]) -> List[str]:
    """Surfaced tensions — where framing and evidence diverge.

    Sources, in order:
      1. ``session._hypothesis_tensions`` (simulate_hypothesis flow)
      2. ``synthesis.tensions[]`` if persisted
      3. ``reasoning_audit_log`` entries with engine='tension_detector'
    """
    candidates: List[Any] = []
    if session.get("_hypothesis_tensions"):
        candidates.extend(session["_hypothesis_tensions"])
    synth = session.get("synthesis") or {}
    if synth.get("tensions"):
        candidates.extend(synth["tensions"])
    if not candidates:
        for e in (session.get("reasoning_audit_log") or []):
            if (e.get("engine") or "").lower() == "tension_detector":
                output = e.get("output") or {}
                tensions = output.get("tensions") or output.get("found") or []
                candidates.extend(tensions)
    out: List[str] = []
    for t in candidates:
        if isinstance(t, dict):
            txt = t.get("description") or t.get("text") or t.get("summary") or ""
        else:
            txt = str(t)
        txt = _strip_tier_markers(txt).strip()
        if txt:
            out.append(txt)
    # Dedup while preserving order.
    seen = set()
    deduped = []
    for x in out:
        if x not in seen:
            deduped.append(x)
            seen.add(x)
    return deduped[:3]


def _format_duration(session: Dict[str, Any]) -> Optional[str]:
    started = session.get("started_at")
    finished = session.get("completed_at") or session.get("updated_at")
    if not (started and finished):
        return None
    try:
        a = datetime.fromisoformat(started.replace("Z", "+00:00"))
        b = datetime.fromisoformat(finished.replace("Z", "+00:00"))
    except (TypeError, ValueError):
        return None
    delta = b - a
    minutes = int(delta.total_seconds() // 60)
    if minutes < 1:
        return "under a minute"
    if minutes == 1:
        return "1 minute"
    if minutes < 90:
        return f"{minutes} minutes"
    hours = minutes // 60
    rem = minutes % 60
    if rem:
        return f"{hours}h {rem}m"
    return f"{hours} hours"


def _framing_one_liner(session: Dict[str, Any]) -> str:
    text = (session.get("intent") or "").strip()
    if not text:
        return "Solva session"
    # Take up to first sentence or 18 words, whichever is shorter.
    first = re.split(r"(?<=[.!?])\s+", text, maxsplit=1)[0]
    words = first.split()
    if len(words) > 18:
        first = " ".join(words[:18]).rstrip(",.;:") + "…"
    return first


def _date_str(session: Dict[str, Any]) -> str:
    when = session.get("completed_at") or session.get("started_at") or ""
    if not when:
        return datetime.utcnow().strftime("%d %B %Y")
    try:
        dt = datetime.fromisoformat(when.replace("Z", "+00:00"))
    except (TypeError, ValueError):
        return datetime.utcnow().strftime("%d %B %Y")
    return dt.strftime("%d %B %Y")


# ---------------------------------------------------------------------------
# Public — context dict (also handy for tests + the frontend if exposed)
# ---------------------------------------------------------------------------
def build_artefact_context(session: Dict[str, Any]) -> Dict[str, Any]:
    """Shape a session row into the dict consumed by the HTML template
    and the DOCX builder. Pure function: no DB calls, no LLM calls."""
    refusal = _is_refusal(session)
    synth = session.get("synthesis") or {}
    claims = synth.get("claims") or []
    scenarios = _scenarios_from_claims(claims)
    diagnosis = _split_paragraphs(synth.get("body") or synth.get("stripped_text") or "")
    recs = _recommendations_for_artefact(session)
    tensions = _tension_items(session)
    sensitivity = _sensitivity_items(session, scenarios)
    submodule = session.get("submodule") or "seek_clarity"
    persona = (session.get("persona") or "").strip() or None

    title = SUBMODULE_LABELS.get(submodule, "Solva session")
    if persona:
        title = f"{title} — {persona}"

    audit_count = len(session.get("reasoning_audit_log") or [])

    ctx = {
        "is_refusal":        refusal,
        "title":             f"Solva · {title}",
        "submodule_label":   SUBMODULE_LABELS.get(submodule, "Solva"),
        "submodule_key":     submodule,
        "persona":           persona,
        "framing_one_line":  _framing_one_liner(session),
        "framing_full":      (session.get("intent") or "").strip(),
        "cluster_label":     session.get("cluster_label") or "",
        "duration_label":    _format_duration(session),
        "date_str":          _date_str(session),
        "diagnosis_paragraphs": diagnosis,
        "scenarios":         scenarios,
        "sensitivity_items": sensitivity,
        "tension_items":     tensions,
        "recommendations":   recs,
        "session_id_short":  (session.get("id") or "")[:8],
        "session_id":        session.get("id") or "",
        "audit_count":       audit_count,
        "status":            session.get("status") or "active",
    }
    if refusal:
        ctx.update(_refusal_overlay(session))
    return ctx


def _refusal_overlay(session: Dict[str, Any]) -> Dict[str, Any]:
    """Replace the 5-section diagnosis with the 4-section refusal anatomy
    described in brief §5.5.

    Sections:
      1. Masthead (kept from main context)
      2. What's missing
      3. What Solva can offer
      4. Recommended next action
    """
    audit = session.get("reasoning_audit_log") or []
    refusal_entries = [e for e in audit if (e.get("engine") or "").lower() == "refusal"]
    last = refusal_entries[-1] if refusal_entries else {}
    output = last.get("output") if isinstance(last, dict) else {}
    output = output or {}

    whats_missing = (
        output.get("missing_evidence")
        or output.get("reason")
        or "Solva does not have enough grounded evidence to weight scenarios honestly for this question."
    )
    candidates_offered = output.get("candidates_for_user") or output.get("candidate_set") or []
    next_actions = output.get("next_actions") or output.get("user_next_steps") or []

    if isinstance(whats_missing, list):
        whats_missing = " ".join(str(x) for x in whats_missing)

    if not candidates_offered:
        # Fall back to whatever candidates the candidate-generation engine
        # produced before refusal triggered.
        for e in audit:
            if (e.get("engine") or "").lower() == "candidate_generation":
                cands = (e.get("output") or {}).get("candidates") or []
                candidates_offered = [
                    (c.get("hypothesis") if isinstance(c, dict) else str(c))
                    for c in cands
                ][:7]
                break

    if not next_actions:
        next_actions = [
            "Pull the source records that would let Solva weight scenarios.",
            "Return for a full synthesis once the evidence gap is closed.",
        ]

    return {
        "diagnosis_paragraphs": [],   # hidden in refusal mode
        "scenarios":            [],   # hidden in refusal mode
        "sensitivity_items":    [],   # hidden in refusal mode
        "tension_items":        [],   # hidden in refusal mode
        "recommendations":      [],   # hidden in refusal mode
        "refusal_whats_missing": _strip_tier_markers(str(whats_missing)),
        "refusal_candidates":    [
            _strip_tier_markers(str(x))
            for x in candidates_offered
            if x
        ],
        "refusal_next_actions":  [
            _strip_tier_markers(str(x))
            for x in next_actions
            if x
        ],
    }


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def _render_html(session: Dict[str, Any]) -> str:
    ctx = build_artefact_context(session)
    if ctx.get("is_refusal"):
        tmpl = _env.get_template("solva_refusal_artefact.html")
    else:
        tmpl = _env.get_template("solva_artefact.html")
    return tmpl.render(**ctx)


def build_pdf(session: Dict[str, Any]) -> bytes:
    """Render the artefact (or refusal artefact) to PDF bytes.

    Caller is responsible for HTTP wrapping. Raises if WeasyPrint is
    not installed (``ImportError``) or template render fails.
    """
    from weasyprint import HTML  # imported lazily so the router stays
                                 # importable without WeasyPrint installed.
    html = _render_html(session)
    pdf_bytes: bytes = HTML(string=html).write_pdf()
    return pdf_bytes


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def build_docx(session: Dict[str, Any]) -> bytes:
    """Render the artefact to DOCX bytes. Probability bars become a 1-row
    inline table; callouts are bordered text boxes. Same Georgia /
    Calibri typography as the PDF / web."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    ctx = build_artefact_context(session)
    is_refusal = ctx["is_refusal"]
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    INK    = RGBColor(0x2A, 0x1B, 0x1D)
    DEEP   = RGBColor(0x5A, 0x4A, 0x4D)
    MUTED  = RGBColor(0x6B, 0x6B, 0x6B)
    ACCENT = RGBColor(0xC2, 0x5A, 0x38)

    def _set_cell_bg(cell, hex_color: str) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _add_para(text, *, font="Calibri", size=11, color=INK, bold=False,
                  italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text or "")
        run.font.name = font
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = color
        p.paragraph_format.space_after = Pt(space_after)
        return p

    # Masthead
    _add_para(
        ctx["submodule_label"] + (f" · {ctx['persona']}" if ctx.get("persona") else ""),
        font="Calibri", size=9, color=MUTED, space_after=2,
    )
    _add_para(ctx["framing_one_line"], font="Georgia", size=22, color=INK, bold=True, space_after=8)
    meta_bits = [ctx["date_str"]]
    if ctx.get("duration_label"):
        meta_bits.append(ctx["duration_label"])
    if ctx.get("cluster_label"):
        meta_bits.append(ctx["cluster_label"])
    _add_para(" · ".join(meta_bits), font="Calibri", size=10, color=DEEP, space_after=14)

    if is_refusal:
        _add_para("WHAT'S MISSING", font="Georgia", size=10, color=ACCENT, italic=True, space_after=6)
        _add_para(ctx.get("refusal_whats_missing") or "", font="Georgia", size=12, color=INK, space_after=14)

        if ctx.get("refusal_candidates"):
            _add_para("WHAT SOLVA CAN OFFER", font="Georgia", size=10, color=ACCENT, italic=True, space_after=4)
            _add_para(
                "Here are the framings worth examining, without weighting.",
                font="Calibri", size=10, color=DEEP, italic=True, space_after=6,
            )
            for cand in ctx["refusal_candidates"]:
                doc.add_paragraph(cand, style="List Bullet")
            _add_para("", space_after=8)

        if ctx.get("refusal_next_actions"):
            _add_para("RECOMMENDED NEXT ACTION", font="Georgia", size=10, color=ACCENT, italic=True, space_after=4)
            for step in ctx["refusal_next_actions"]:
                doc.add_paragraph(step, style="List Number")

    else:
        # Primary diagnosis
        _add_para("PRIMARY DIAGNOSIS", font="Georgia", size=10, color=ACCENT, italic=True, space_after=6)
        for para in (ctx["diagnosis_paragraphs"] or ["Diagnosis not yet available."]):
            _add_para(para, font="Georgia", size=12, color=INK, space_after=8)

        # Scenarios — bars as 1-row inline tables
        if ctx["scenarios"]:
            _add_para("", space_after=4)
            _add_para("SCENARIOS", font="Georgia", size=10, color=ACCENT, italic=True, space_after=6)
            for s in ctx["scenarios"]:
                head = doc.add_paragraph()
                head.paragraph_format.space_after = Pt(2)
                r1 = head.add_run(s["label"])
                r1.font.name = "Georgia"
                r1.font.size = Pt(11)
                r1.bold = True
                r1.font.color.rgb = INK
                head.add_run("    ")
                r2 = head.add_run(f"{s['pct']}% ({s['low']}–{s['high']}%)")
                r2.font.name = "Calibri"
                r2.font.size = Pt(10)
                r2.font.color.rgb = DEEP

                if s["desc"]:
                    _add_para(s["desc"], font="Georgia", size=10, color=DEEP, italic=True, space_after=4)

                # 1x100 visual bar — three regions: pre-pct INK, CI overlay (left of pt? we use 3 cells)
                # Simple 3-col bar: [0..low: rule], [low..pct: ink], [pct..100: rule]
                low = s["low"]
                high = s["high"]
                pct = s["pct"]
                widths_pct = [
                    max(0, low),                 # left pad (faint)
                    max(0, pct - low),           # ink (point estimate body)
                    max(0, high - pct),          # CI extension to high (semi)
                    max(0, 100 - high),          # right pad
                ]
                # Avoid 0-width cells (Word renders them oddly).
                widths_pct = [max(0.5, w) if w > 0 else 0 for w in widths_pct]
                shown = [w for w in widths_pct if w > 0]
                cols = max(1, len(shown))
                tbl = doc.add_table(rows=1, cols=cols)
                tbl.autofit = False
                row = tbl.rows[0]
                row.height = Cm(0.42)
                colors = []
                for w_idx, w in enumerate(widths_pct):
                    if w == 0:
                        continue
                    if w_idx == 1:           # ink (point estimate)
                        colors.append(("2A1B1D", w))
                    elif w_idx == 2:         # CI overlay
                        colors.append(("8B6F71", w))   # muted accent
                    else:                    # rule
                        colors.append(("D5C9B6", w))
                # Apply widths and fills
                total_cm = 14.0
                for i, (hexc, w) in enumerate(colors):
                    cell = row.cells[i]
                    cell.width = Cm(total_cm * (w / 100.0))
                    _set_cell_bg(cell, hexc)
                    # Empty paragraph keeps the cell tall
                    cell.text = ""
                _add_para("", space_after=6)

        if ctx.get("sensitivity_items"):
            _add_para("WHAT WOULD CHANGE THIS READ", font="Georgia", size=10,
                      color=ACCENT, italic=True, space_after=4)
            tbl = doc.add_table(rows=1, cols=1)
            cell = tbl.rows[0].cells[0]
            _set_cell_bg(cell, "F5EFE6")
            inner = cell.paragraphs[0]
            inner.paragraph_format.space_after = Pt(0)
            cell.text = ""
            for it in ctx["sensitivity_items"]:
                p = cell.add_paragraph(it)
                for run in p.runs:
                    run.font.name = "Georgia"
                    run.font.size = Pt(11)
                    run.font.color.rgb = INK
            _add_para("", space_after=8)

        if ctx.get("tension_items"):
            _add_para("WHERE YOUR FRAMING AND THE EVIDENCE DIVERGE", font="Georgia",
                      size=10, color=ACCENT, italic=True, space_after=4)
            tbl = doc.add_table(rows=1, cols=1)
            cell = tbl.rows[0].cells[0]
            _set_cell_bg(cell, "E8DCC8")
            cell.text = ""
            for it in ctx["tension_items"]:
                p = cell.add_paragraph(it)
                for run in p.runs:
                    run.font.name = "Georgia"
                    run.font.size = Pt(11)
                    run.font.color.rgb = INK
            _add_para("", space_after=8)

        if ctx.get("recommendations"):
            _add_para("RECOMMENDATIONS", font="Georgia", size=10, color=ACCENT, italic=True, space_after=6)
            for r in ctx["recommendations"]:
                p = doc.add_paragraph(style="List Number")
                run_h = p.add_run(f"{r['heading']} — ")
                run_h.font.name = "Georgia"
                run_h.font.size = Pt(11)
                run_h.bold = True
                run_h.font.color.rgb = INK
                run_b = p.add_run(r["body"])
                run_b.font.name = "Georgia"
                run_b.font.size = Pt(11)
                run_b.font.color.rgb = INK

    # Footer
    _add_para("", space_after=12)
    _add_para(
        f"Session {ctx['session_id_short']} · Audit log {ctx['audit_count']} entries · "
        f"{ctx['date_str']}",
        font="Calibri", size=9, color=MUTED, space_after=2,
    )
    _add_para("Generated by AKKI · Solva v3", font="Calibri", size=9, color=MUTED)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = [
    "build_pdf",
    "build_docx",
    "build_artefact_context",
    "SUBMODULE_LABELS",
]
