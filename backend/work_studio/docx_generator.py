"""
DOCX generator — programmatic python-docx; no Jinja, no template
files. Mimics the WeDeliver concept-note signature:
  - cover block (title, subtitle, framework spine, host org, version, date, audience)
  - two-tier numbered headings (§1, §1.1, §1.2, ...)
  - distinct heading + body fonts (Georgia headings, Calibri body)
  - structured tables at high fidelity; bulleted lists at low fidelity
  - section-end soft separator line
"""
from __future__ import annotations
from io import BytesIO
from typing import List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .brief import Brief, BriefSection, BriefTable, FIDELITY_HIGH, FIDELITY_LOW

# Brand palette ------------------------------------------------------------
INK   = RGBColor(0x1A, 0x1D, 0x20)
OXBLD = RGBColor(0x7A, 0x2E, 0x2E)
MUTED = RGBColor(0x6F, 0x71, 0x77)

BODY_FONT    = "Calibri"
HEADING_FONT = "Georgia"
KICKER_FONT  = "Calibri"


def _set_run(run, *, font=BODY_FONT, size_pt=11, bold=False, italic=False, colour=INK):
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = colour
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)
    rfonts.set(qn('w:cs'), font)


def _add_para(doc, text, *, font=BODY_FONT, size_pt=11, bold=False, italic=False,
              colour=INK, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
              space_before=0, all_caps=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text.upper() if all_caps else text)
    _set_run(run, font=font, size_pt=size_pt, bold=bold, italic=italic, colour=colour)
    if all_caps:
        run.font.all_caps = True
    return p


def _add_heading_1(doc, kicker_num: str, title: str):
    """§N kicker on its own line, then heading 1 text."""
    _add_para(doc, kicker_num, font=KICKER_FONT, size_pt=10, bold=True,
              colour=OXBLD, space_before=18, space_after=2, all_caps=True)
    _add_para(doc, title, font=HEADING_FONT, size_pt=18, bold=True,
              colour=INK, space_after=8)


def _add_heading_2(doc, kicker_num: str, title: str):
    _add_para(doc, kicker_num, font=KICKER_FONT, size_pt=9, bold=True,
              colour=OXBLD, space_before=10, space_after=2, all_caps=True)
    _add_para(doc, title, font=HEADING_FONT, size_pt=13, bold=True,
              colour=INK, space_after=6)


def _add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    _set_run(run, size_pt=11, colour=INK)


def _add_table(doc, t: BriefTable, *, fidelity: str):
    """High-fidelity: real docx table with header row shaded.
    Low-fidelity:    bullet list per row."""
    if fidelity == FIDELITY_LOW:
        _add_para(doc, t.title.upper(), font=KICKER_FONT, size_pt=9, bold=True,
                  colour=MUTED, space_before=6, space_after=2)
        for row in t.rows:
            _add_bullet(doc, " · ".join(str(c) for c in row if str(c).strip()))
        return

    _add_para(doc, t.title.upper(), font=KICKER_FONT, size_pt=9, bold=True,
              colour=OXBLD, space_before=8, space_after=4, all_caps=True)

    table = doc.add_table(rows=1 + len(t.rows), cols=len(t.headers))
    table.style = "Table Grid"
    # Header
    for i, h in enumerate(t.headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h.upper())
        _set_run(run, font=KICKER_FONT, size_pt=9, bold=True, colour=RGBColor(0xFF, 0xFF, 0xFF))
        # Shade header
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0E0E0E')
        tcPr.append(shd)
    # Body
    for r_idx, row in enumerate(t.rows):
        for c_idx, val in enumerate(row):
            if c_idx >= len(t.headers):
                break
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _set_run(run, size_pt=10, colour=INK)
    doc.add_paragraph()  # trailing spacer


def _add_separator(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'C8C8C8')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section(doc, idx: int, sec: BriefSection, *, fidelity: str):
    _add_heading_1(doc, f"§ {idx}", sec.title)
    if sec.kicker:
        _add_para(doc, sec.kicker, font=KICKER_FONT, size_pt=9, bold=True,
                  colour=MUTED, all_caps=True, space_after=4)
    for sub_idx, p in enumerate(sec.body_paragraphs):
        # Two-tier headings only when paragraph starts with "Heading: ..." marker.
        # Otherwise emit as plain body.
        _add_para(doc, p, size_pt=11, colour=INK, space_after=8)
    for b in sec.bullets:
        _add_bullet(doc, b)
    for t in sec.tables:
        _add_table(doc, t, fidelity=fidelity)
    _add_separator(doc)


def render_docx(brief: Brief) -> bytes:
    doc = Document()
    # 1.0" margins
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # ------------- Cover block --------------
    _add_para(doc, brief.document_type.upper(),
              font=KICKER_FONT, size_pt=10, bold=True, colour=OXBLD,
              align=WD_ALIGN_PARAGRAPH.LEFT, all_caps=True, space_after=4)
    if brief.subtitle:
        _add_para(doc, brief.subtitle,
                  font=KICKER_FONT, size_pt=11, italic=True, colour=MUTED,
                  space_after=18)
    _add_para(doc, brief.title, font=HEADING_FONT, size_pt=28, bold=True,
              colour=INK, space_after=10)
    if brief.cover_lead_paragraph:
        _add_para(doc, brief.cover_lead_paragraph,
                  font=BODY_FONT, size_pt=12, colour=INK, space_after=18)
    if brief.framework_spine:
        _add_para(doc, brief.framework_spine,
                  font=KICKER_FONT, size_pt=12, bold=True, colour=OXBLD,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_after=18, all_caps=True)
    if brief.host_org_line:
        _add_para(doc, "Hosted by", font=KICKER_FONT, size_pt=9,
                  bold=True, colour=MUTED, all_caps=True, space_after=2)
        _add_para(doc, brief.host_org_line, font=HEADING_FONT, size_pt=12,
                  colour=INK, space_after=14)
    meta_bits = []
    if brief.version: meta_bits.append(f"Document version: {brief.version}")
    if brief.date_text: meta_bits.append(brief.date_text)
    if brief.audience: meta_bits.append(f"Audience: {brief.audience}")
    if meta_bits:
        _add_para(doc, "  ·  ".join(meta_bits), font=BODY_FONT, size_pt=10,
                  colour=MUTED, space_after=14)
    # Page break before sections at high fidelity, double separator at low.
    if brief.fidelity == FIDELITY_HIGH:
        doc.add_page_break()
    else:
        _add_separator(doc); _add_separator(doc)

    # ------------- Sections --------------
    for i, sec in enumerate(brief.sections, start=1):
        _add_section(doc, i, sec, fidelity=brief.fidelity)

    # ------------- Closing --------------
    if brief.closing_recap:
        _add_para(doc, "RECAP", font=KICKER_FONT, size_pt=9, bold=True,
                  colour=OXBLD, all_caps=True, space_before=14, space_after=4)
        _add_para(doc, brief.closing_recap, size_pt=11, colour=INK, space_after=14)
    if brief.closing_brand_line:
        _add_para(doc, brief.closing_brand_line, font=KICKER_FONT, size_pt=10,
                  colour=MUTED, italic=True, all_caps=True, space_after=0)

    # STUDIO sprint — W-19: Synisense audit footer line. Only rendered
    # when Brief.audit_summary is set (preserves determinism for fixtures
    # that don't supply one).
    if brief.audit_summary:
        _add_para(
            doc, brief.audit_summary, font="Courier New", size_pt=9,
            colour=MUTED, italic=True, space_before=18, space_after=0,
        )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
